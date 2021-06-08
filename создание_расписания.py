import docx
import os
from random import randint
from datetime import timedelta, datetime


def ворлд(now, schedule_territories):
    now_day = datetime(now.year, now.month, now.day)
    name_file = 'расписание_охраны.docx'
    document = docx.Document()
    document.save(name_file)
    doc = docx.Document(name_file)
    doc.add_heading('Докладная записка', 1)
    doc.add_heading('Докладываю Вам, что с {} по {} обходы территории склада осуществлялись:'.format(now_day,now_day + timedelta(days=1)), 3)
    
    for territory in schedule_territories:
        doc.add_heading('{}__________:'.format(territory), 3)

        for schedule in schedule_territories[territory]:
            doc1 = doc.add_heading('Убыл:{}'.format(schedule), 4)
            doc1.add_run('   Прибыл:____________')
    doc.save(name_file)


def get_a_schedule_rounded_to_5(set_time, base=5):
    territories = ['ГСМ', 'Cтоянка', 'Ангар']
    schedule_territories = {}

    for territory in territories:
        schedule = []
        crawl_time = set_time

        while crawl_time.hour != 7 and crawl_time.hour != 6:
            time_between_crawl = randint(90, 120)
            time_between_crawl = base * round(time_between_crawl/base)
            after_the_crawl = timedelta(minutes=time_between_crawl)
            crawl_time = crawl_time+after_the_crawl
            schedule.append(crawl_time)
        schedule_territories.update({territory: schedule})
    return schedule_territories


def get_date_and_time():
    now = datetime.now()
    set_time = datetime(now.year, now.month, now.day, hour=8, minute=0, second=0, microsecond=0)
    return now, set_time


if __name__ == '__main__':
    now, set_time = get_date_and_time()
    schedule_territories = get_a_schedule_rounded_to_5(set_time)
    ворлд(now, schedule_territories)
    